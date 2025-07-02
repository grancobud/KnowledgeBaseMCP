#!/usr/bin/env python3
"""
DOCX Writer - Word belgesi oluşturma modülü
"""

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)

class DocxWriter:
    """DOCX belgesi oluşturma sınıfı"""
    
    def __init__(self):
        logger.info("DocxWriter başlatıldı")
    
    async def create_document(self, content: str, file_path: str, title: str = None, 
                            style_config: Dict[str, Any] = None) -> str:
        """
        DOCX belgesi oluştur
        
        Args:
            content: Belge içeriği
            file_path: Kayıt edilecek dosya yolu
            title: Belge başlığı
            style_config: Stil yapılandırması
            
        Returns:
            str: İşlem sonucu mesajı
        """
        try:
            # Yeni doküman oluştur
            doc = Document()
            
            # Varsayılan stil yapılandırması
            if style_config is None:
                style_config = {
                    "font_name": "Calibri",
                    "font_size": 11,
                    "line_spacing": 1.15,
                    "add_page_numbers": True,
                    "margins": 1.0  # inch
                }
            
            # Sayfa ayarları
            self._setup_page_settings(doc, style_config)
            
            # Başlık ekle
            if title:
                self._add_title(doc, title, style_config)
            
            # İçeriği parse et ve ekle
            self._add_content(doc, content, style_config)
            
            # Sayfa numaraları ekle
            if style_config.get("add_page_numbers", True):
                self._add_page_numbers(doc)
            
            # Dosyayı kaydet
            doc.save(file_path)
            
            logger.info(f"DOCX belgesi oluşturuldu: {file_path}")
            return f"DOCX belgesi başarıyla oluşturuldu: {file_path}"
            
        except Exception as e:
            error_msg = f"DOCX oluşturma hatası: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def _setup_page_settings(self, doc: Document, style_config: Dict[str, Any]):
        """Sayfa ayarlarını yapılandır"""
        sections = doc.sections
        for section in sections:
            # Sayfa kenar boşlukları
            margin = Inches(style_config.get("margins", 1.0))
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin
    
    def _add_title(self, doc: Document, title: str, style_config: Dict[str, Any]):
        """Başlık ekle"""
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Başlık stilini ayarla
        run = title_paragraph.runs[0]
        run.font.name = style_config.get("font_name", "Calibri")
        run.font.size = Pt(16)
        run.bold = True
    
    def _add_content(self, doc: Document, content: str, style_config: Dict[str, Any]):
        """İçeriği parse et ve ekle"""
        lines = content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Boş satır - yeni paragraf
                if current_paragraph is not None:
                    current_paragraph = None
                continue
            
            # Başlık kontrolü (=== ile çevrili)
            if line.startswith('===') and line.endswith('==='):
                heading_text = line.replace('=', '').strip()
                if heading_text:
                    heading = doc.add_heading(heading_text, level=2)
                    self._style_heading(heading, style_config)
                current_paragraph = None
                continue
            
            # Alt başlık kontrolü (-- ile başlayan)
            if line.startswith('--') or line.startswith('##'):
                heading_text = line.replace('-', '').replace('#', '').strip()
                if heading_text:
                    heading = doc.add_heading(heading_text, level=3)
                    self._style_heading(heading, style_config)
                current_paragraph = None
                continue
            
            # Liste öğesi kontrolü
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                list_text = line[1:].strip()
                if list_text:
                    p = doc.add_paragraph(list_text, style='List Bullet')
                    self._style_paragraph(p, style_config)
                current_paragraph = None
                continue
            
            # Numaralı liste kontrolü
            if len(line) > 2 and line[0].isdigit() and line[1] == '.':
                list_text = line[2:].strip()
                if list_text:
                    p = doc.add_paragraph(list_text, style='List Number')
                    self._style_paragraph(p, style_config)
                current_paragraph = None
                continue
            
            # Normal paragraf
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
                self._style_paragraph(current_paragraph, style_config)
            
            # Bold text kontrolü (**text**)
            if '**' in line:
                self._add_formatted_text(current_paragraph, line)
            else:
                current_paragraph.add_run(line + ' ')
    
    def _add_formatted_text(self, paragraph, text: str):
        """Formatted text ekle (bold, italic vb.)"""
        parts = text.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Normal text
                if part:
                    paragraph.add_run(part)
            else:
                # Bold text
                if part:
                    run = paragraph.add_run(part)
                    run.bold = True
    
    def _style_paragraph(self, paragraph, style_config: Dict[str, Any]):
        """Paragraf stilini ayarla"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = style_config.get("line_spacing", 1.15)
        paragraph_format.space_after = Pt(6)
        
        for run in paragraph.runs:
            run.font.name = style_config.get("font_name", "Calibri")
            run.font.size = Pt(style_config.get("font_size", 11))
    
    def _style_heading(self, heading, style_config: Dict[str, Any]):
        """Başlık stilini ayarla"""
        for run in heading.runs:
            run.font.name = style_config.get("font_name", "Calibri")
            run.bold = True
    
    def _add_page_numbers(self, doc: Document):
        """Sayfa numaraları ekle"""
        try:
            # Footer'a sayfa numarası ekle
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Sayfa numarası field kodu ekle
            run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
            
            # Field code for page numbers
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
        except Exception as e:
            logger.warning(f"Sayfa numarası ekleme hatası: {e}")

    async def create_structured_report(self, data: Dict[str, Any], file_path: str) -> str:
        """
        Yapılandırılmış rapor oluştur
        
        Args:
            data: Rapor verisi
            file_path: Kayıt edilecek dosya yolu
            
        Returns:
            str: İşlem sonucu
        """
        try:
            doc = Document()
            
            # Başlık
            if 'title' in data:
                title = doc.add_heading(data['title'], level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Özet bölümü
            if 'summary' in data:
                doc.add_heading('EXECUTIVE SUMMARY', level=2)
                summary_p = doc.add_paragraph(data['summary'])
                summary_p.paragraph_format.space_after = Pt(12)
            
            # Bölümler
            if 'sections' in data:
                for section in data['sections']:
                    # Bölüm başlığı
                    doc.add_heading(section.get('title', 'Başlıksız Bölüm'), level=2)
                    
                    # Bölüm içeriği
                    if 'content' in section:
                        if isinstance(section['content'], list):
                            for item in section['content']:
                                doc.add_paragraph(str(item), style='List Bullet')
                        else:
                            doc.add_paragraph(str(section['content']))
                    
                    # Alt bölümler
                    if 'subsections' in section:
                        for subsection in section['subsections']:
                            doc.add_heading(subsection.get('title', 'Alt Başlık'), level=3)
                            if 'content' in subsection:
                                doc.add_paragraph(str(subsection['content']))
            
            # Dosyayı kaydet
            doc.save(file_path)
            
            return f"Yapılandırılmış rapor oluşturuldu: {file_path}"
            
        except Exception as e:
            error_msg = f"Yapılandırılmış rapor oluşturma hatası: {str(e)}"
            logger.error(error_msg)
            return error_msg
